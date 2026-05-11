# services/ingest/app/services/table_extractor.py
#
# Tabellen-Extraktor: Liest Tabellen aus verschiedenen Dokumentformaten
# und wandelt sie in semantisch sinnvolle Prosa-Chunks um.
#
# Unterstützte Formate:
#   Schritt 1: DOCX        → python-docx  (exakte XML-Struktur)
#   Schritt 2: PDF         → Docling      (beste PDF-Tabellenerkennung)
#   Schritt 3: XLSX / ODS  → pandas       (Spreadsheets als Tabellen-Chunks)
#
# Alle Pfade münden in dieselbe TableObject-Struktur die einheitlich
# in Prosa-Sätze und Chunks umgewandelt wird.

from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

import structlog

logger = structlog.get_logger()


# ─────────────────────────────────────────────────────────────────────────────
# Datenklassen
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class TableCell:
    """Eine einzelne Tabellenzelle."""
    text:      str
    row:       int
    col:       int
    row_span:  int = 1
    col_span:  int = 1
    is_header: bool = False


@dataclass
class TableObject:
    """
    Eine extrahierte Tabelle – formatunabhängige Zwischendarstellung.
    Wird von allen Extraktoren erzeugt und einheitlich verarbeitet.
    """
    headers:     list[str]               # Spaltenköpfe (erste Zeile)
    rows:        list[list[str]]         # Datenzeilen als String-Listen
    caption:     str = ""                # Tabellenüberschrift (optional)
    section_ref: str = ""                # Abschnitt in dem die Tabelle steht
    source_fmt:  str = ""                # docx | pdf | xlsx | ods | html
    table_index: int = 0                 # Reihenfolge im Dokument

    @property
    def is_empty(self) -> bool:
        return not self.rows or not any(
            any(c.strip() for c in row) for row in self.rows
        )

    @property
    def col_count(self) -> int:
        return len(self.headers)

    @property
    def row_count(self) -> int:
        return len(self.rows)


@dataclass
class TableChunk:
    """Eine Tabellenzeile als Prosa-Text für das Chunking."""
    text:        str
    row_index:   int
    table_index: int
    caption:     str = ""
    section_ref: str = ""
    source_fmt:  str = ""


# ─────────────────────────────────────────────────────────────────────────────
# Tabelle → Prosa Konvertierung
# ─────────────────────────────────────────────────────────────────────────────

class TableToProsa:
    """
    Wandelt TableObject-Zeilen in semantisch sinnvolle Prosa-Sätze um.

    Strategie 1 (Standard): Schlüssel-Wert-Satz
      "Normtyp: Pflicht. Formulierung: muss / ist zu. Tag: MUST."

    Strategie 2 (bei 2 Spalten): Aussagesatz
      "Pflicht bedeutet: muss / ist zu."

    Strategie 3 (Fallback): Aufzählung
      "Pflicht | muss / ist zu | MUST"
    """

    def convert(
        self,
        table: TableObject,
        include_caption: bool = True,
    ) -> list[TableChunk]:
        """Konvertiert alle Zeilen einer Tabelle in TableChunk-Objekte."""
        if table.is_empty:
            return []

        chunks = []

        # Präambel aus Überschrift + Spaltenkontext
        preamble = ""
        if include_caption and table.caption:
            preamble = f"{table.caption}: "

        for i, row in enumerate(table.rows):
            # Leere Zeilen überspringen
            if not any(c.strip() for c in row):
                continue

            text = self._row_to_prosa(
                row=row,
                headers=table.headers,
                preamble=preamble,
            )

            if text.strip():
                chunks.append(TableChunk(
                    text=text,
                    row_index=i,
                    table_index=table.table_index,
                    caption=table.caption,
                    section_ref=table.section_ref,
                    source_fmt=table.source_fmt,
                ))

        return chunks

    def _row_to_prosa(
        self,
        row:      list[str],
        headers:  list[str],
        preamble: str = "",
    ) -> str:
        """Eine Zeile → Prosa-Satz."""
        # Zellen bereinigen
        cells = [c.strip() for c in row]
        hdrs  = [h.strip() for h in headers]

        # Leere Zellen ausblenden
        pairs = [
            (h, c) for h, c in zip(hdrs, cells)
            if c and h
        ]
        # Zellen ohne Header ergänzen
        for i, c in enumerate(cells):
            if i >= len(hdrs) and c:
                pairs.append((f"Spalte {i+1}", c))

        if not pairs:
            return ""

        # Strategie 2: Nur 2 Spalten → Aussagesatz
        if len(pairs) == 2:
            k, v = pairs[0][1], pairs[1][1]
            h0, h1 = pairs[0][0], pairs[1][0]
            return f"{preamble}{h0} \"{k}\": {h1} ist \"{v}\"."

        # Strategie 1: Schlüssel-Wert-Aufzählung
        kv_parts = ". ".join(f"{h}: {v}" for h, v in pairs)
        return f"{preamble}{kv_parts}."

    def table_summary(self, table: TableObject) -> str:
        """
        Erzeugt einen einzeiligen Zusammenfassungs-Chunk für die gesamte Tabelle.
        Wird als Parent-Chunk verwendet.
        """
        info = []
        if table.caption:
            info.append(table.caption)
        if table.headers:
            info.append(f"Spalten: {', '.join(table.headers)}")
        info.append(f"{table.row_count} Einträge")
        return ". ".join(info) + "."


# ─────────────────────────────────────────────────────────────────────────────
# Schritt 1: DOCX-Tabellen-Extraktor
# ─────────────────────────────────────────────────────────────────────────────

class DocxTableExtractor:
    """
    Extrahiert Tabellen aus DOCX-Dateien via python-docx.
    DOCX-Tabellen sind exakt im XML-Modell definiert –
    keine Heuristik nötig.
    """

    def extract(
        self,
        content: bytes,
        filename: str = "",
    ) -> list[TableObject]:
        """
        Liest alle Tabellen aus einem DOCX-Dokument.

        Returns:
            Liste von TableObject (eine pro Tabelle im Dokument)
        """
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx nicht installiert – pip install python-docx")
            return []

        try:
            doc = Document(io.BytesIO(content))
        except Exception as e:
            logger.error("DOCX konnte nicht geöffnet werden", error=str(e))
            return []

        tables = []

        # Abschnitts-Kontext aufbauen: welche Überschrift steht vor der Tabelle?
        section_map = self._build_section_map(doc)

        for i, tbl in enumerate(doc.tables):
            table_obj = self._extract_table(tbl, i, section_map)
            if table_obj and not table_obj.is_empty:
                tables.append(table_obj)

        logger.info(
            "DOCX Tabellen extrahiert",
            filename=filename,
            count=len(tables),
        )
        return tables

    def _extract_table(
        self,
        tbl,
        index: int,
        section_map: dict,
    ) -> Optional[TableObject]:
        """Extrahiert eine einzelne python-docx Table."""
        if not tbl.rows:
            return None

        rows_raw = []
        for row in tbl.rows:
            row_cells = []
            for cell in row.cells:
                # Merged cells haben denselben Text mehrfach –
                # Duplikate innerhalb einer Zeile entfernen
                text = cell.text.strip().replace("\n", " ")
                row_cells.append(text)
            # Aufeinanderfolgende identische Zellen (Merges) deduplizieren
            deduped = []
            prev = None
            for c in row_cells:
                if c != prev:
                    deduped.append(c)
                prev = c
            rows_raw.append(deduped)

        if not rows_raw:
            return None

        # Erste Zeile als Header interpretieren wenn sie sich von den
        # Folgezeilen unterscheidet (übliche DOCX-Tabellen-Konvention)
        headers = rows_raw[0]
        data_rows = rows_raw[1:] if len(rows_raw) > 1 else []

        # Abschnitts-Kontext
        section_ref = section_map.get(index, "")

        return TableObject(
            headers=headers,
            rows=data_rows,
            caption="",
            section_ref=section_ref,
            source_fmt="docx",
            table_index=index,
        )

    def _build_section_map(self, doc) -> dict:
        """
        Baut eine Map: Tabellen-Index → nächste vorherige Überschrift.
        Hilft beim Kontextualisieren der Tabelle.
        """
        section_map = {}
        table_idx = 0
        last_heading = ""

        # Paragraphen und Tabellen in Dokumentreihenfolge durchlaufen
        from docx.oxml.ns import qn
        body = doc.element.body

        for child in body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag == 'p':
                # Überschrift?
                style = child.find(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pStyle'
                )
                if style is not None:
                    style_val = style.get(
                        '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', ''
                    )
                    if 'Heading' in style_val or 'berschrift' in style_val:
                        # Text der Überschrift
                        texts = child.itertext()
                        last_heading = ''.join(texts).strip()
            elif tag == 'tbl':
                section_map[table_idx] = last_heading
                table_idx += 1

        return section_map


# ─────────────────────────────────────────────────────────────────────────────
# Schritt 2: PDF-Tabellen-Extraktor (Docling)
# ─────────────────────────────────────────────────────────────────────────────

class PdfTableExtractor:
    """
    Extrahiert Tabellen aus PDF-Dateien via Docling.
    Docling liefert strukturierte Tabellen-Objekte mit
    Zeile/Spalte/Span-Information.
    """

    def extract(
        self,
        content: bytes,
        filename: str = "",
    ) -> list[TableObject]:
        """
        Liest alle Tabellen aus einem PDF via Docling.

        Returns:
            Liste von TableObject
        """
        try:
            from docling.document_converter import DocumentConverter
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.pipeline_options import PdfPipelineOptions
        except ImportError:
            logger.warning("Docling nicht installiert – pip install docling")
            return []

        try:
            # Temporäre Datei – Docling braucht einen Dateipfad
            import tempfile, os
            with tempfile.NamedTemporaryFile(
                suffix=".pdf", delete=False
            ) as tmp:
                tmp.write(content)
                tmp_path = tmp.name

            pipeline_opts = PdfPipelineOptions()
            pipeline_opts.do_table_structure = True    # Tabellenerkennung aktiv
            pipeline_opts.do_ocr              = False  # OCR nur wenn nötig

            converter = DocumentConverter()
            result    = converter.convert(tmp_path)
            doc       = result.document

            os.unlink(tmp_path)

        except Exception as e:
            logger.error("PDF-Docling-Parsing fehlgeschlagen", error=str(e))
            return []

        tables = []
        for i, table in enumerate(doc.tables):
            table_obj = self._convert_docling_table(table, i)
            if table_obj and not table_obj.is_empty:
                tables.append(table_obj)

        logger.info(
            "PDF Tabellen extrahiert (Docling)",
            filename=filename,
            count=len(tables),
        )
        return tables

    def _convert_docling_table(self, table, index: int) -> Optional[TableObject]:
        """Konvertiert ein Docling-Table-Objekt in ein TableObject."""
        try:
            # Docling liefert grid: list[list[TableCell]]
            grid = table.data.grid
            if not grid:
                return None

            # Erste Zeile als Header
            headers = []
            for cell in grid[0]:
                headers.append(cell.text.strip() if cell.text else "")

            # Restliche Zeilen als Daten
            rows = []
            for row in grid[1:]:
                row_data = []
                for cell in row:
                    row_data.append(cell.text.strip() if cell.text else "")
                if any(c for c in row_data):
                    rows.append(row_data)

            # Beschriftung aus Docling-Kontext
            caption = ""
            if hasattr(table, 'caption') and table.caption:
                caption = str(table.caption).strip()

            return TableObject(
                headers=headers,
                rows=rows,
                caption=caption,
                section_ref="",
                source_fmt="pdf",
                table_index=index,
            )
        except Exception as e:
            logger.warning("Docling-Tabelle konnte nicht konvertiert werden",
                           error=str(e))
            return None


# ─────────────────────────────────────────────────────────────────────────────
# Schritt 3: Spreadsheet-Tabellen-Extraktor (XLSX / ODS)
# ─────────────────────────────────────────────────────────────────────────────

class SpreadsheetTableExtractor:
    """
    Extrahiert Tabellen aus XLSX- und ODS-Dateien via pandas.
    Jedes Tabellenblatt wird als eigene Tabelle behandelt.
    """

    def extract(
        self,
        content:  bytes,
        filename: str = "",
    ) -> list[TableObject]:
        """
        Liest alle Tabellenblätter aus XLSX oder ODS.

        Returns:
            Liste von TableObject (eine pro Sheet)
        """
        try:
            import pandas as pd
        except ImportError:
            logger.warning("pandas nicht installiert – pip install pandas openpyxl odfpy")
            return []

        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        try:
            if ext in ("xlsx", "xls"):
                engine = "openpyxl" if ext == "xlsx" else "xlrd"
                sheets = pd.read_excel(
                    io.BytesIO(content),
                    sheet_name=None,   # alle Sheets
                    engine=engine,
                    dtype=str,         # alles als String
                    na_filter=False,   # NaN → leer lassen
                )
            elif ext == "ods":
                sheets = pd.read_excel(
                    io.BytesIO(content),
                    sheet_name=None,
                    engine="odf",
                    dtype=str,
                    na_filter=False,
                )
            else:
                logger.warning("Unbekanntes Spreadsheet-Format", ext=ext)
                return []

        except Exception as e:
            logger.error("Spreadsheet konnte nicht geöffnet werden",
                         filename=filename, error=str(e))
            return []

        tables = []
        for sheet_idx, (sheet_name, df) in enumerate(sheets.items()):
            table_obj = self._df_to_table(df, sheet_name, sheet_idx)
            if table_obj and not table_obj.is_empty:
                tables.append(table_obj)

        logger.info(
            "Spreadsheet Tabellen extrahiert",
            filename=filename,
            sheets=len(sheets),
            non_empty=len(tables),
        )
        return tables

    def _df_to_table(
        self,
        df,
        sheet_name: str,
        index:      int,
    ) -> Optional[TableObject]:
        """Konvertiert einen pandas DataFrame in ein TableObject."""
        import pandas as pd

        # Leere DataFrames überspringen
        if df.empty:
            return None

        # Leere Zeilen und Spalten entfernen
        df = df.dropna(how="all").dropna(axis=1, how="all")
        if df.empty:
            return None

        # Erste Zeile als Header wenn sie keine Zahlen enthält
        first_row = [str(v).strip() for v in df.iloc[0]]
        if all(not self._is_numeric(v) for v in first_row if v):
            headers  = first_row
            data_rows = df.iloc[1:]
        else:
            # Spalten-Index als Header verwenden
            headers   = [str(c) for c in df.columns]
            data_rows = df

        rows = []
        for _, row in data_rows.iterrows():
            row_data = [str(v).strip() for v in row]
            if any(v for v in row_data):
                rows.append(row_data)

        return TableObject(
            headers=headers,
            rows=rows,
            caption=sheet_name,
            section_ref="",
            source_fmt="xlsx" if "xls" in str(index) else "spreadsheet",
            table_index=index,
        )

    def _is_numeric(self, s: str) -> bool:
        """Prüft ob ein String eine Zahl ist."""
        try:
            float(s.replace(",", "."))
            return True
        except (ValueError, AttributeError):
            return False


# ─────────────────────────────────────────────────────────────────────────────
# Zentraler Tabellen-Extraktor (Dispatcher)
# ─────────────────────────────────────────────────────────────────────────────

class TableExtractor:
    """
    Zentraler Dispatcher – wählt den richtigen Extraktor
    je nach Dateiformat.
    """

    def __init__(self):
        self.docx_extractor  = DocxTableExtractor()
        self.pdf_extractor   = PdfTableExtractor()
        self.sheet_extractor = SpreadsheetTableExtractor()
        self.converter       = TableToProsa()

    def extract_tables(
        self,
        content:  bytes,
        filename: str,
    ) -> list[TableObject]:
        """
        Wählt den passenden Extraktor und gibt TableObjects zurück.
        """
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""

        if ext == "docx":
            return self.docx_extractor.extract(content, filename)
        elif ext == "pdf":
            return self.pdf_extractor.extract(content, filename)
        elif ext in ("xlsx", "xls", "ods"):
            return self.sheet_extractor.extract(content, filename)
        else:
            # Kein dedizierter Tabellen-Extraktor – Tika-Fallback
            return []

    def tables_to_text(
        self,
        tables: list[TableObject],
        include_summary: bool = True,
    ) -> str:
        """
        Wandelt alle Tabellen in annotierten Fließtext um.
        Wird in parser.py verwendet um Tika-Text zu ergänzen.

        Format:
          [TABELLE 1: Sheet-Name / Caption]
          Zeile als Prosa-Satz.
          Zeile als Prosa-Satz.
          [/TABELLE 1]
        """
        parts = []
        for table in tables:
            if table.is_empty:
                continue

            # Tabellen-Header
            header_info = table.caption or f"Tabelle {table.table_index + 1}"
            if table.section_ref:
                header_info += f" (aus: {table.section_ref})"
            parts.append(f"\n[TABELLE {table.table_index + 1}: {header_info}]")

            # Zusammenfassung
            if include_summary:
                parts.append(self.converter.table_summary(table))

            # Zeilen als Prosa
            chunks = self.converter.convert(table, include_caption=False)
            for chunk in chunks:
                parts.append(chunk.text)

            parts.append(f"[/TABELLE {table.table_index + 1}]\n")

        return "\n".join(parts)

    def has_extractor(self, filename: str) -> bool:
        """Gibt True zurück wenn ein dedizierter Extraktor vorhanden ist."""
        ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
        return ext in ("docx", "pdf", "xlsx", "xls", "ods")
